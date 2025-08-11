# app.py - Clean version with no duplicates
from docx import Document
import io
import requests
from flask import Flask, request, jsonify, send_file
import tempfile
import os
import time
from datetime import datetime

app = Flask(__name__)

@app.route('/', methods=['GET'])
def root():
    """Root endpoint"""
    return jsonify({
        'message': 'Document Processing Service',
        'status': 'ready',
        'endpoints': {
            'health': '/health',
            'process': '/process-word-document'
        }
    })

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'message': 'Word document processor with signature image preservation',
        'timestamp': datetime.now().isoformat(),
        'preserves': ['signature images', 'all formatting', 'logos', 'layout'],
        'template_id': '1jnQRnscY9chDJZMPY8EI5TsQScMRNF05'
    })

@app.route('/process-word-document', methods=['POST'])
def process_word_document():
    """
    Process Word document while preserving ALL images and formatting
    """
    try:
        print("=== Document Processing Started ===")
        
        # Get data from n8n request
        data = request.json
        seller_name = data.get('seller_name')
        
        if not seller_name:
            return jsonify({'error': 'seller_name is required'}), 400
        
        print(f"Processing document for: {seller_name}")
        
        # Your template file ID from Google Drive
        template_file_id = '1jnQRnscY9chDJZMPY8EI5TsQScMRNF05'
        download_url = f'https://drive.google.com/uc?export=download&id={template_file_id}'
        
        print(f"Downloading template from: {download_url}")
        
        # Download the template file
        response = requests.get(download_url)
        if response.status_code != 200:
            print(f"Download failed with status: {response.status_code}")
            return jsonify({'error': f'Could not download template. Status: {response.status_code}'}), 400
        
        print(f"Template downloaded successfully. Size: {len(response.content)} bytes")
        
        # Load the Word document from the downloaded content
        doc_stream = io.BytesIO(response.content)
        doc = Document(doc_stream)
        
        print("Document loaded into python-docx successfully")
        
        # Prepare replacements
        replacements = {
            '{{SellerName}}': seller_name,
            '{{PropertyAddress}}': data.get('property_address', 'TBD'),
            '{{ContractDate}}': data.get('contract_date', 'TBD'),
            '{{NotaryDate}}': data.get('notary_date', datetime.now().strftime("%B %d, %Y"))
        }
        
        print(f"Replacements to make: {replacements}")
        
        replacement_count = 0
        
        # Replace text in paragraphs (images are preserved automatically)
        print("Processing paragraphs...")
        for i, paragraph in enumerate(doc.paragraphs):
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    print(f"Found {placeholder} in paragraph {i}")
                    # Replace in runs to preserve formatting
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            replacement_count += 1
                            print(f"  Replaced {placeholder} with {value}")
        
        # Replace text in tables (images in tables are also preserved)
        print("Processing tables...")
        for table_num, table in enumerate(doc.tables):
            for row_num, row in enumerate(table.rows):
                for cell_num, cell in enumerate(row.cells):
                    for para_num, paragraph in enumerate(cell.paragraphs):
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                print(f"Found {placeholder} in table {table_num}, row {row_num}, cell {cell_num}")
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)
                                        replacement_count += 1
                                        print(f"  Replaced {placeholder} with {value} in table")
        
        # Replace text in headers and footers (preserves images there too)
        print("Processing headers and footers...")
        for section_num, section in enumerate(doc.sections):
            # Headers
            for para_num, paragraph in enumerate(section.header.paragraphs):
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        print(f"Found {placeholder} in header {section_num}")
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
                                replacement_count += 1
                                print(f"  Replaced {placeholder} with {value} in header")
            
            # Footers
            for para_num, paragraph in enumerate(section.footer.paragraphs):
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        print(f"Found {placeholder} in footer {section_num}")
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
                                replacement_count += 1
                                print(f"  Replaced {placeholder} with {value} in footer")
        
        print(f"Total replacements made: {replacement_count}")
        
        # Save to temporary file
        output_filename = f"Processed_Affidavit_{seller_name.replace(' ', '_').replace(',', '')}_{int(time.time())}.docx"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        print(f"Document saved to: {temp_file.name}")
        print("✅ All signature images preserved!")
        
        # Return the processed file
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Error processing document: {str(e)}")
        return jsonify({
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

if __name__ == '__main__':
    PORT = 5001
    
    print("🚀 Starting Document Processing Service...")
    print("📄 Template preserves: Signature images, formatting, logos")
    print(f"🌐 Service will be available at: http://localhost:{PORT}")
    print(f"📋 Health check: http://localhost:{PORT}/health")
    print(f"🔧 Processing endpoint: http://localhost:{PORT}/process-word-document")
    print(f"🌍 Also available at: http://192.168.100.18:{PORT}")
    
    app.run(host='0.0.0.0', port=PORT, debug=True)